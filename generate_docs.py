import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell margins (padding) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_color(doc, text, level, space_before=120, space_after=60):
    """Adds a heading with styled colors and margins."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before / 20)
    heading.paragraph_format.space_after = Pt(space_after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors
    # H1 = Slate Blue, H2 = Teal, H3 = Charcoal
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(13, 148, 136) # Teal 600
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        run.bold = True
    return heading

def main():
    doc = docx.Document()
    
    # Set standard page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set default style to Arial, 11pt, charcoal text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    
    # Set default paragraph spacing
    p_format = style.paragraph_format
    p_format.space_after = Pt(6)
    p_format.line_spacing = 1.15

    # ----------------------------------------------------
    # TITLE PAGE
    # ----------------------------------------------------
    # Add vertical spacing
    for _ in range(3):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("🏠 StaySphere")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("A Modern PG & Hostel Management System")
    sub_run.font.size = Pt(16)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    for _ in range(6):
        doc.add_paragraph()
        
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.line_spacing = 1.5
    
    info_run1 = info_p.add_run("Complete Technical Documentation & User Manual\n")
    info_run1.font.size = Pt(12)
    info_run1.font.bold = True
    
    info_run2 = info_p.add_run(
        "Prepared For: Academic Evaluation & Project Submission\n"
        "GitHub Repository: https://github.com/yashwanterla-cmyk/stay-sphere\n"
        "Webapp Deployment Link: https://stay-sphere.vercel.app\n"
        "Document Version: 1.0\n"
        "Date: July 2026\n"
    )
    info_run2.font.size = Pt(10.5)
    info_run2.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()

    # ----------------------------------------------------
    # TABLE OF CONTENTS / EXECUTIVE SUMMARY
    # ----------------------------------------------------
    add_heading_with_color(doc, "1. Executive Summary", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "StaySphere is an advanced, full-stack property management system designed specifically for Paying Guest (PG) "
        "accommodations, hostels, and rental properties. In the modern hospitality and rental market, property owners "
        "struggle with a variety of operational pain points, including manual bed allocation tracking, inconsistent tenant "
        "onboarding, tracking recurring monthly rental invoices, processing security check-ins, logging external visitors, "
        "and coordinating maintenance complaints. StaySphere addresses these pain points by offering a single, "
        "role-based workspace where owners, managers, staff, and tenants can collaborate smoothly."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "By integrating security checks, automated invoicing, visual bed allocation management, a digital agreement signing pad, "
        "and real-time financial reporting, StaySphere streamlines property workflows. This document provides a exhaustive look "
        "into the platform's features, architectural blueprint, database structure, API specification, visual pages walkthrough, "
        "and operational setup guides."
    )
    
    # ----------------------------------------------------
    # SYSTEM FEATURES
    # ----------------------------------------------------
    add_heading_with_color(doc, "2. Key Modules & Features", level=1)
    
    modules = [
        ("🔐 Authentication & Role-Based Access Control", 
         "StaySphere uses secure JWT-based authentication. Users are categorized into four distinct roles: "
         "Super Admin (full platform visibility), Owner (manages their own properties, rooms, and tenant operations), "
         "Staff (handles daily operation tasks like visitor logs, attendance, and maintenance tickets), and "
         "Tenant (accesses invoice history, signs lease agreements, and files maintenance issues)."),
        
        ("🏢 Property & Bed Management", 
         "Owners can register multiple properties. When adding rooms, the system automatically instantiates bed units "
         "(e.g., Room 101 with Double occupancy generates 101-A and 101-B). Beds are color-coded in real-time "
         "indicating their operational state: Vacant, Occupied, or under Maintenance."),
        
        ("👥 Tenant Registration & KYC", 
         "Facilitates smooth digital onboarding of tenants. Staff can enter emergency contacts, guardian credentials, "
         "and upload national identification documents (e.g. Aadhaar, Passport) stored securely in Cloudinary."),
        
        ("💰 Rent Billing & Invoicing", 
         "Invoices are auto-generated monthly for active tenants. The platform integrates a simulated Razorpay payment gateway "
         "allowing tenants to make mock online payments, which verify payment signatures and instantly issue printable receipts."),
        
        ("🔧 Maintenance Ticketing System", 
         "Tenants can file complaints categorized by issue type (Plumbing, Electrical, Housekeeping) and priority (Low, Medium, High). "
         "Owners or staff can assign internal staff members to issues and trace the ticket timeline from 'Pending' to 'Resolved'."),
        
        ("👁️ Security & Visitor Logbook", 
         "Logs all incoming property visitors. Security staff can record a visitor's name, contact, purpose of visit, check-in timestamp, "
         "and check-out timestamp, maintaining a robust audit trail for property security."),
        
        ("📅 Attendance Logging", 
         "Staff and tenants can log daily check-ins and check-outs, assisting administrators in tracking employee working hours "
         "and monitoring tenant occupancy patterns."),
        
        ("💸 Expense Tracking", 
         "Allows owners to log operational expenses (e.g., electricity bills, water, salary payouts, repair expenses). "
         "Expenses are grouped and charted against monthly rental collections to evaluate net profit margins."),
        
        ("📄 Digital Rental Agreements", 
         "Integrates a dynamic signing pad utilizing HTML5 canvas. Tenants can review terms, sign using touch or mouse gestures, "
         "and compile signed agreements with custom digital hash tags stored in the database.")
    ]
    
    for title, desc in modules:
        add_heading_with_color(doc, title, level=2)
        doc.add_paragraph(desc)
        
    doc.add_page_break()

    # ----------------------------------------------------
    # SYSTEM ARCHITECTURE
    # ----------------------------------------------------
    add_heading_with_color(doc, "3. System Architecture", level=1)
    
    doc.add_paragraph(
        "StaySphere uses a decoupled Client-Server architecture to guarantee high performance, modularity, and ease of deployment. "
        "The web client communicates with the central backend server through a RESTful API, secured by bearer tokens."
    )
    
    add_heading_with_color(doc, "3.1 Technology Stack Matrix", level=2)
    
    # Table for Tech Stack
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technologies'
    hdr_cells[2].text = 'Role & Importance'
    
    # Style Header Row
    for cell in hdr_cells:
        set_cell_background(cell, '1E293B') # Dark Slate
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    tech_data = [
        ("Frontend Client", "React 19, TypeScript, Vite, Tailwind CSS 3.4", "Provides a type-safe, responsive Single Page Application (SPA). Uses React Context for state, Framer Motion for animations, and Recharts for interactive dashboards."),
        ("Backend API Server", "FastAPI (Python 3.10+), Pydantic v2", "Offers high-performance, asynchronous routing with automated OpenAPI documentation (Swagger UI). Validates payloads using strict Pydantic schemas."),
        ("Database & ORM", "SQLAlchemy 2.0, SQLite (Dev), PostgreSQL (Prod)", "Ensures robust ORM mapping, relational integrity, migrations, and support for scalable production database backends."),
        ("Security & Auth", "python-jose, passlib[bcrypt]", "Handles password hashing using salted bcrypt algorithms and issues JWT access tokens with 7-day lifespans."),
        ("Cloud Integrations", "Razorpay SDK, Cloudinary python SDK", "Simulates digital payment order verification and securely hosts tenant KYC photo assets and signature files.")
    ]
    
    for layer, tech, role in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = role
        # Add padding & thin borders
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_background(cell, 'F8FAFC') # Soft gray
            
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # DATABASE DESIGN
    # ----------------------------------------------------
    add_heading_with_color(doc, "4. Database Design & Relational Schema", level=1)
    
    doc.add_paragraph(
        "StaySphere uses a highly structured relational database schema. The backend enforces strict relational constraints "
        "to manage properties, rooms, beds, tenants, and invoices without data anomalies. For example, deleting a room cascade "
        "deletes its beds, but tenant records are preserved or unlinked appropriately."
    )
    
    # Table for DB Schema
    db_table = doc.add_table(rows=1, cols=4)
    db_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = db_table.rows[0].cells
    hdr[0].text = 'Model Name'
    hdr[1].text = 'Database Table'
    hdr[2].text = 'Primary Keys / Key Fields'
    hdr[3].text = 'Key Relationships'
    
    for cell in hdr:
        set_cell_background(cell, '0F766E') # Teal 700
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    db_models = [
        ("User", "users", "id, email, password_hash, role, status", "One-to-One with Tenant, One-to-Many with Attendance log"),
        ("Property", "properties", "id, name, address, owner_id", "One-to-Many with Room, Expense, and Visitor"),
        ("Room", "rooms", "id, room_number, price_per_bed, capacity", "Many-to-One with Property, One-to-Many with Bed"),
        ("Bed", "beds", "id, bed_number, status, room_id", "Many-to-One with Room, One-to-One with Tenant"),
        ("Tenant", "tenants", "id, user_id, bed_id, status, guardian_name", "One-to-One with User, One-to-Many with Invoice & Agreement"),
        ("Invoice", "invoices", "id, amount, status, tenant_id", "Many-to-One with Tenant"),
        ("MaintenanceRequest", "maintenance_requests", "id, title, priority, status, tenant_id", "Many-to-One with Tenant & Property"),
        ("Visitor", "visitors", "id, name, phone, check_in_time, property_id", "Many-to-One with Property"),
        ("Attendance", "attendances", "id, user_id, check_in, check_out", "Many-to-One with User"),
        ("Expense", "expenses", "id, category, amount, property_id", "Many-to-One with Property"),
        ("Agreement", "agreements", "id, status, tenant_id, signed_at", "Many-to-One with Tenant")
    ]
    
    for model, tbl, keys, rels in db_models:
        row_cells = db_table.add_row().cells
        row_cells[0].text = model
        row_cells[1].text = tbl
        row_cells[2].text = keys
        row_cells[3].text = rels
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_background(cell, 'F0FDFA') # Mint tint

    doc.add_page_break()

    # ----------------------------------------------------
    # API ENDPOINTS DOCUMENTATION
    # ----------------------------------------------------
    add_heading_with_color(doc, "5. RESTful API Reference", level=1)
    
    doc.add_paragraph(
        "The backend exposes a secure API prefix under `/api/v1`. Authentication is mandated via a JWT Bearer header. "
        "Below is an overview of the critical endpoints deployed on the server:"
    )
    
    # Table for APIs
    api_table = doc.add_table(rows=1, cols=4)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = api_table.rows[0].cells
    hdr[0].text = 'HTTP Method'
    hdr[1].text = 'Endpoint Path'
    hdr[2].text = 'Required Role'
    hdr[3].text = 'Description / Action'
    
    for cell in hdr:
        set_cell_background(cell, '475569') # Slate 600
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    apis = [
        ("POST", "/api/v1/auth/signup", "Any", "Register user account (Owner or Tenant)"),
        ("POST", "/api/v1/auth/login", "Any", "Authenticate credentials & obtain Bearer JWT"),
        ("GET", "/api/v1/auth/me", "Authenticated", "Retrieve current user profile data"),
        ("GET", "/api/v1/properties", "Owner / Admin", "List all properties owned by user"),
        ("POST", "/api/v1/rooms/property/{id}", "Owner", "Create a room and auto-populate bed units"),
        ("POST", "/api/v1/tenants/register", "Owner / Staff", "Register a tenant and assign a vacant bed"),
        ("GET", "/api/v1/rent/invoices", "Tenant / Owner", "List invoices (filtered by tenant)"),
        ("POST", "/api/v1/rent/pay/{id}", "Tenant", "Initiate simulated Razorpay order payment"),
        ("POST", "/api/v1/maintenance", "Tenant / Staff", "File maintenance issues with priority level"),
        ("PUT", "/api/v1/maintenance/{id}", "Owner / Staff", "Update ticket status (in_progress, resolved)"),
        ("POST", "/api/v1/visitors", "Staff / Owner", "Log entry of external visitor on property"),
        ("POST", "/api/v1/agreements/{id}/sign", "Tenant", "Submit canvas Base64 signature for lease")
    ]
    
    for method, path, role, desc in apis:
        row_cells = api_table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = path
        row_cells[2].text = role
        row_cells[3].text = desc
        
        # Color HTTP Method
        if method == "POST":
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(22, 163, 74) # Green
        elif method == "GET":
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(37, 99, 235) # Blue
        elif method == "PUT":
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(217, 119, 6) # Amber
        
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_background(cell, 'F8FAFC')
            
    doc.add_page_break()

    # ----------------------------------------------------
    # VISUAL PAGES WALKTHROUGH
    # ----------------------------------------------------
    add_heading_with_color(doc, "6. User Interface Walkthrough", level=1)
    
    ui_walkthroughs = [
        ("6.1 Dashboard Workspace",
         "Upon logging in, the owner sees an overview dashboard. The upper cards display real-time statistics: "
         "total monthly earnings, current occupancy rates, pending maintenance tickets, and unpaid tenant invoices. "
         "Below the cards, interactive charts display monthly income vs. expense. A live activity timeline "
         "displays visitors checking in, notices being published, and tenant signatures completed."),
        
        ("6.2 Visual Bed Allocation",
         "The Room Management interface lets owners see a grid of all rooms. Selecting a room displays "
         "individual beds (A, B, C). Beds colored in Green are vacant and available for allotment. Clicking on a vacant bed "
         "opens the 'Onboard Tenant' registration form. Beds colored in Red represent active occupants, showing "
         "the tenant's name and lease start dates. Yellow beds indicate a bed under maintenance."),
        
        ("6.3 Invoice Panel & Razorpay Simulation",
         "Tenants can navigate to their personal billing page, listing all historical invoices. Unpaid invoices have "
         "a 'Pay Now' action. When triggered, a custom-designed checkout screen modal appears. Upon completing the simulated "
         "checkout, the backend handles verification and updates the invoice state, issuing a PDF download receipt for the tenant."),
        
        ("6.4 Digital Lease Pad",
         "When a tenant onboard, a lease agreement is generated automatically. The tenant opens the agreement page, reads "
         "the rental terms, and uses a drawing pad. The canvas records mouse/touch coordinates and saves "
         "the signature as a base64 image string. On clicking 'Submit', the agreement updates to 'Active', capturing the legal timestamp.")
    ]
    
    for title, walk in ui_walkthroughs:
        add_heading_with_color(doc, title, level=2)
        doc.add_paragraph(walk)

    # ----------------------------------------------------
    # SETUP AND DEPLOYMENT
    # ----------------------------------------------------
    add_heading_with_color(doc, "7. Setup & Installation Guide", level=1)
    
    doc.add_paragraph(
        "To deploy StaySphere locally, follow the steps detailed below. Ensure you have Node.js 18+ and Python 3.10+ installed."
    )
    
    add_heading_with_color(doc, "7.1 Backend API Server Setup", level=2)
    doc.add_paragraph(
        "1. Open a terminal and navigate to the backend folder:\n"
        "   cd backend\n\n"
        "2. Create a virtual environment:\n"
        "   python -m venv .venv\n"
        "   Windows: .venv\\Scripts\\activate\n"
        "   macOS/Linux: source .venv/bin/activate\n\n"
        "3. Install python packages:\n"
        "   pip install -r requirements.txt\n\n"
        "4. Configure environment variables in `.env` (API keys for Cloudinary and Razorpay ID).\n\n"
        "5. Launch the FastAPI server:\n"
        "   python -m uvicorn app.main:app --reload --port 8000"
    )
    
    add_heading_with_color(doc, "7.2 Frontend React Setup", level=2)
    doc.add_paragraph(
        "1. Open a new terminal in the frontend directory:\n"
        "   cd frontend\n\n"
        "2. Install packages:\n"
        "   npm install\n\n"
        "3. Configure env parameters if needed (VITE_API_URL defaults to localhost:8000).\n\n"
        "4. Launch development dev-server:\n"
        "   npm run dev\n\n"
        "5. Access the application in your browser at: http://localhost:5173"
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # PROMPT ENGINEERING SUMMARY
    # ----------------------------------------------------
    add_heading_with_color(doc, "8. AI Prompt Engineering & Code Generation", level=1)
    doc.add_paragraph(
        "A key requirement of this project was utilizing AI-assisted software engineering. The development team leveraged "
        "advanced LLMs to generate complex algorithms, mock payment gateways, and configure the database relational model. "
        "A dedicated `prompt_library.md` file has been added to the root directory, documenting the precise system prompts, "
        "role playing settings, and structural guidelines used to construct the API routes and UI components.\n\n"
        "By cataloging these prompt iterations, the project remains highly reproducible, letting other developers "
        "extend modules like automated notice publishing or biometric check-ins using identical context criteria."
    )
    
    # Save Document
    doc_path = "Project_Documentation.docx"
    doc.save(doc_path)
    print(f"[SUCCESS] Documentation created successfully: {os.path.abspath(doc_path)}")

if __name__ == "__main__":
    main()
